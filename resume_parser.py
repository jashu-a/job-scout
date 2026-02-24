"""
Resume parser - extracts text from DOCX or PDF resumes.
"""

from pathlib import Path

try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False


def extract_resume_text(resume_path: str) -> str:
    """Extract text from a DOCX or PDF resume."""
    path = Path(resume_path)

    if not path.exists():
        raise FileNotFoundError(f"Resume not found: {resume_path}")

    suffix = path.suffix.lower()

    # ── DOCX ──
    if suffix == ".docx":
        return _extract_docx(path)

    # ── PDF ──
    if suffix == ".pdf":
        return _extract_pdf(path)

    # ── Plain text fallback ──
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8")

    raise ValueError(f"Unsupported resume format: {suffix}. Use .docx, .pdf, or .txt")


def _extract_docx(path: Path) -> str:
    """Extract text from DOCX, including text inside tables (for multi-column resumes)."""
    if not HAS_PYTHON_DOCX:
        raise RuntimeError("python-docx is required for DOCX parsing: pip install python-docx")

    doc = DocxDocument(str(path))
    parts = []

    # Extract from main body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract from tables (many modern resumes use tables for column layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text and text not in parts:  # Avoid duplicates from nested tables
                        parts.append(text)

    if not parts:
        raise RuntimeError(f"Could not extract any text from {path}")

    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF. Tries pdfplumber first, falls back to pypdf."""
    # Try pdfplumber (better layout extraction)
    if HAS_PDFPLUMBER:
        try:
            text = ""
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text.strip()
        except Exception:
            pass

    # Fallback to pypdf
    if HAS_PYPDF:
        try:
            reader = PdfReader(str(path))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if text.strip():
                return text.strip()
        except Exception:
            pass

    raise RuntimeError(
        f"Could not extract text from {path}. "
        "Install pdfplumber or pypdf: pip install pdfplumber pypdf"
    )
