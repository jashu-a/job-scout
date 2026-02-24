"""
Document generator - creates tailored resume by cloning the original DOCX
and replacing content in-place, preserving all formatting.

Also generates a cover letter as a clean new DOCX.
"""

import re
import shutil
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ═══════════════════════════════════════════════════════════════════════════════
# TAILORED RESUME — Clone & Replace Strategy
# ═══════════════════════════════════════════════════════════════════════════════

def _build_replacement_map(resume_data: dict) -> dict:
    """
    Build a flat mapping of { original_text_fragment: replacement_text }
    from the AI-generated tailored resume data.

    The AI returns structured data with original and rewritten bullets.
    We match on experience entries by company name, then swap bullets.
    """
    return resume_data


def _replace_run_text(run, old_text: str, new_text: str) -> bool:
    """Replace text in a single run, preserving all run formatting."""
    if old_text in run.text:
        run.text = run.text.replace(old_text, new_text)
        return True
    return False


def _find_and_replace_in_paragraph(paragraph, old_text: str, new_text: str) -> bool:
    """
    Replace text across a paragraph's runs while preserving formatting.

    Handles the common case where a single logical string is split across
    multiple XML runs (due to spell-check, formatting changes, etc.).
    """
    # Fast path: check if the full text even contains what we're looking for
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # Try simple per-run replacement first
    for run in paragraph.runs:
        if _replace_run_text(run, old_text, new_text):
            return True

    # If text is split across runs, do a cross-run replacement
    # Strategy: put all new text in the first run that contains part of old_text,
    # clear the others that were part of the match
    runs = paragraph.runs
    if not runs:
        return False

    # Build a map of character positions to runs
    char_positions = []  # [(run_index, char_in_run)]
    for ri, run in enumerate(runs):
        for ci in range(len(run.text)):
            char_positions.append((ri, ci))

    start_idx = full_text.find(old_text)
    if start_idx == -1:
        return False

    end_idx = start_idx + len(old_text)

    # Find which runs are involved
    start_run_idx = char_positions[start_idx][0] if start_idx < len(char_positions) else None
    end_run_idx = char_positions[end_idx - 1][0] if end_idx - 1 < len(char_positions) else None

    if start_run_idx is None or end_run_idx is None:
        return False

    # Replace: put new text where old text started, clear the rest
    involved_runs = list(range(start_run_idx, end_run_idx + 1))

    # Reconstruct: text before match in first run + new text + text after match in last run
    first_run = runs[involved_runs[0]]
    last_run = runs[involved_runs[-1]]

    # Calculate offsets within the first and last runs
    chars_before_first_run = sum(len(runs[i].text) for i in range(involved_runs[0]))
    offset_in_first = start_idx - chars_before_first_run

    chars_before_last_run = sum(len(runs[i].text) for i in range(involved_runs[-1]))
    offset_in_last = end_idx - chars_before_last_run

    prefix = first_run.text[:offset_in_first]
    suffix = last_run.text[offset_in_last:]

    # Set first run's text to prefix + new_text + (suffix if same run)
    if len(involved_runs) == 1:
        first_run.text = prefix + new_text + suffix
    else:
        first_run.text = prefix + new_text
        last_run.text = suffix
        # Clear middle runs
        for ri in involved_runs[1:-1]:
            runs[ri].text = ""

    return True


def _replace_in_document(doc: Document, old_text: str, new_text: str) -> int:
    """
    Find and replace text throughout a DOCX document (paragraphs + table cells).
    Returns count of replacements made.
    """
    count = 0

    # Body paragraphs
    for para in doc.paragraphs:
        if _find_and_replace_in_paragraph(para, old_text, new_text):
            count += 1

    # Table cells (for column-layout resumes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _find_and_replace_in_paragraph(para, old_text, new_text):
                        count += 1

    # Headers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    if _find_and_replace_in_paragraph(para, old_text, new_text):
                        count += 1

    # Footers
    for section in doc.sections:
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    if _find_and_replace_in_paragraph(para, old_text, new_text):
                        count += 1

    return count


def create_tailored_resume(
    original_docx_path: str,
    resume_data: dict,
    job_title: str,
    job_company: str,
    output_path: str,
) -> str:
    """
    Create a tailored resume by cloning the original DOCX and performing
    targeted text replacements based on AI-generated content.

    The original document's formatting, styles, layout, fonts, columns,
    tables, margins, and all visual properties are preserved exactly.

    Args:
        original_docx_path: Path to the candidate's original resume.docx
        resume_data:        Dict from matcher.generate_tailored_resume()
                            Contains replacement_pairs and optionally a new summary.
        job_title:          Target job title (for logging)
        job_company:        Target company (for logging)
        output_path:        Where to save the tailored .docx

    Returns:
        Path to the created file
    """
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    # Step 1: Copy the original file byte-for-byte
    shutil.copy2(original_docx_path, output_path)

    # Step 2: Open the copy and perform replacements
    doc = Document(output_path)
    total_replacements = 0

    # Apply bullet/text replacements from AI
    replacement_pairs = resume_data.get("replacement_pairs", [])
    for pair in replacement_pairs:
        old = pair.get("original", "").strip()
        new = pair.get("tailored", "").strip()
        if old and new and old != new:
            count = _replace_in_document(doc, old, new)
            total_replacements += count
            if count == 0:
                # Try a fuzzy match — sometimes whitespace or punctuation differs slightly
                # Try matching just the first 60 chars
                if len(old) > 60:
                    truncated = old[:60]
                    for para in doc.paragraphs:
                        if truncated in para.text:
                            _find_and_replace_in_paragraph(para, para.text, new)
                            total_replacements += 1
                            break
                    else:
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for para in cell.paragraphs:
                                        if truncated in para.text:
                                            _find_and_replace_in_paragraph(para, para.text, new)
                                            total_replacements += 1
                                            break

    # Apply skills reorder if provided
    skills_replacement = resume_data.get("skills_replacement", {})
    if skills_replacement:
        old_skills = skills_replacement.get("original", "").strip()
        new_skills = skills_replacement.get("tailored", "").strip()
        if old_skills and new_skills:
            count = _replace_in_document(doc, old_skills, new_skills)
            total_replacements += count

    # Apply summary replacement if provided
    summary_replacement = resume_data.get("summary_replacement", {})
    if summary_replacement:
        old_summary = summary_replacement.get("original", "").strip()
        new_summary = summary_replacement.get("tailored", "").strip()
        if old_summary and new_summary:
            count = _replace_in_document(doc, old_summary, new_summary)
            total_replacements += count

    # Save
    doc.save(output_path)

    print(f"       📊 Made {total_replacements} text replacements in cloned resume")
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# COVER LETTER — New document (no template needed)
# ═══════════════════════════════════════════════════════════════════════════════

def create_cover_letter(
    cover_letter_data: dict,
    candidate_name: str,
    contact_info: str,
    job_title: str,
    job_company: str,
    output_path: str,
) -> str:
    """
    Create a cover letter DOCX.

    Args:
        cover_letter_data: Dict from matcher.generate_cover_letter()
        candidate_name:    From the resume data
        contact_info:      From the resume data
        job_title:         Target job title
        job_company:       Target company name
        output_path:       Where to save the .docx

    Returns:
        Path to the created file
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # -- Header: Candidate info --
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(candidate_name)
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    if contact_info:
        contact_para = doc.add_paragraph(contact_info)
        contact_para.paragraph_format.space_after = Pt(12)
        for run in contact_para.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # -- Date --
    from datetime import date
    date_para = doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    date_para.paragraph_format.space_after = Pt(12)

    # -- Addressee --
    doc.add_paragraph(f"RE: {job_title} at {job_company}")
    doc.add_paragraph()

    # -- Body --
    letter_text = cover_letter_data.get("cover_letter", "")
    paragraphs = letter_text.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if para_text:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(8)

    # -- Closing --
    doc.add_paragraph()
    closing = doc.add_paragraph("Best regards,")
    closing.paragraph_format.space_after = Pt(4)
    sig = doc.add_paragraph()
    sig_run = sig.add_run(candidate_name)
    sig_run.bold = True

    # Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
